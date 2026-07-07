"""File-type aware text extractor.

Supported: PDF, DOCX, XLSX, CSV, TXT/MD, HTML. Each extractor returns
either a plain string (prose) or a list of (sheet, rows) tuples
(tabular). The caller branches on the return shape.

Resource caps (`MAX_PAGES`, `MAX_ROWS`, `MAX_CHARS`) bound extraction
so a single pathological file (huge PDF, million-row sheet) cannot
monopolize memory/CPU or produce an unbounded number of chunks.

Security note (C-3)
------------------

Every text result passes through :func:`sanitize_for_context` before
it is returned. This removes the prompt-injection fence delimiters
the orchestrator uses (``<<<CONTEXT_START>>>`` /
``<<<CONTEXT_END>>>``) and a small list of well-known injection
phrases ("ignore previous instructions", "system:", etc.). The
sanitizer is **defense in depth** — the orchestrator already tells
the LLM to treat retrieved chunks as untrusted — but a malicious
document that uses the same fence delimiters could otherwise
trivially escape the untrusted zone.
"""
from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import List, Tuple

from app.core.config import settings
from app.core.logging import get_logger

log = get_logger(__name__)

# Resource caps — defend against decompression / row-count DoS.
MAX_PAGES = 500
MAX_ROWS = 50_000
MAX_CHARS = 2_000_000  # ~ a few hundred pages of text; chunker will split it
MAX_HTML_CHARS = 2_000_000


# --- Prompt-injection sanitization (C-3) ---------------------------------

# Fence delimiters the orchestrator uses to wrap untrusted retrieved
# context. A document that contains the closing delimiter can break
# out of the untrusted zone in the assembled prompt. Strip both.
_FENCE_RE = re.compile(r"<<<\s*/?\s*CONTEXT_(?:START|END)\s*>>>", re.IGNORECASE)

# Lines that look like a chat role prefix — a classic injection
# technique ("system: do the following ..."). Match the line *start*
# only so a normal "System: foo" reference inside prose survives.
_ROLE_PREFIX_RE = re.compile(
    r"(?im)^\s*(system|assistant|user|tool|function)\s*:\s*"
)

# A small list of well-known "jailbreak" prefix phrases. We are
# intentionally conservative: matching the literal phrase catches
# the lazy variants without false-positiving on legitimate prose.
# Long-tail variants ("disregard the prior context", "forget your
# instructions") are caught by the LLM-judge that scores the
# retrieval set; the regex here is a cheap first pass.
_INJECTION_PHRASES = (
    "ignore previous instructions",
    "ignore the above",
    "ignore all previous",
    "ignore your instructions",
    "disregard previous instructions",
    "disregard your instructions",
    "forget your instructions",
    "forget everything above",
    "you are now",
    "new instructions:",
    "system prompt:",
)
_INJECTION_RE = re.compile(
    "|".join(re.escape(p) for p in _INJECTION_PHRASES),
    re.IGNORECASE,
)


def sanitize_for_context(text: str) -> str:
    """Strip prompt-injection patterns from extracted text.

    C-3 — every extractor funnels its output through here before
    returning. The function is **idempotent** (running it twice
    yields the same result) and **lossy on purpose** — a sentence
    that reads "ignore previous instructions" is more dangerous to
    keep than useful to read.

    What it strips:

    1. The orchestrator's own fence delimiters
       (``<<<CONTEXT_START>>>`` / ``<<<CONTEXT_END>>>``), so a
       malicious document cannot break out of the untrusted zone.
    2. Lines that look like a chat role prefix
       (``system:`` / ``assistant:`` / ``user:`` / ``tool:`` /
       ``function:``) at the start of a line. A normal mid-sentence
       "the user said" reference is preserved.
    3. A short list of well-known injection phrases (case-insensitive,
       anywhere in the text). The phrase is replaced with a
       neutral token ``[REDACTED-INJECTION]`` so the chunk length
       and embedding position stay stable.

    What it does **not** do:

    * It does not try to detect every variant of every jailbreak
      prompt — that is the LLM-judge's job downstream. The goal
      here is to remove the cheap, structural attacks.
    * It does not HTML-escape the text. The output is still plain
      prose; React renders it as text, not HTML.
    """
    if not text:
        return text
    # 1. Fence delimiters
    text = _FENCE_RE.sub("", text)
    # 2. Role-prefix lines
    text = _ROLE_PREFIX_RE.sub("", text)
    # 3. Known injection phrases
    text = _INJECTION_RE.sub("[REDACTED-INJECTION]", text)
    return text


class ExtractionResult:
    """Holds the output of an extract step. `mode` is 'prose' or 'tabular'."""

    __slots__ = ("mode", "text", "tables", "meta")

    def __init__(
        self,
        mode: str,
        text: str = "",
        tables: List[Tuple[str, List[List[str]]]] | None = None,
        meta: dict | None = None,
    ) -> None:
        self.mode = mode
        self.text = text
        self.tables = tables or []
        self.meta = meta or {}


def _cap_text(text: str) -> str:
    return text[:MAX_CHARS]


def _ext_pdf(path: Path) -> ExtractionResult:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    total = 0
    for i, page in enumerate(reader.pages):
        if i >= MAX_PAGES:
            log.warning("extract.pdf.page_cap", path=str(path), cap=MAX_PAGES)
            break
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            log.warning("extract.pdf.page_failed", page=i, error=str(exc))
            continue
        if total + len(page_text) > MAX_CHARS:
            page_text = page_text[: max(0, MAX_CHARS - total)]
            parts.append(page_text)
            log.warning("extract.pdf.char_cap", path=str(path))
            break
        parts.append(page_text)
        total += len(page_text)
    return ExtractionResult(mode="prose", text="\n\n".join(parts), meta={"pages": len(parts)})


def _ext_docx(path: Path) -> ExtractionResult:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts = []
    total = 0
    for p in doc.paragraphs:
        if not p.text:
            continue
        if total >= MAX_CHARS:
            log.warning("extract.docx.char_cap", path=str(path))
            break
        parts.append(p.text)
        total += len(p.text)
    return ExtractionResult(mode="prose", text="\n\n".join(parts))


def _ext_xlsx(path: Path) -> ExtractionResult:
    from openpyxl import load_workbook

    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    tables: list[Tuple[str, List[List[str]]]] = []
    total_rows = 0
    for ws in wb.worksheets:
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            if total_rows >= MAX_ROWS:
                log.warning("extract.xlsx.row_cap", sheet=ws.title, cap=MAX_ROWS)
                break
            rows.append(["" if c is None else str(c) for c in row])
            total_rows += 1
        if rows:
            tables.append((ws.title, rows))
        if total_rows >= MAX_ROWS:
            break
    return ExtractionResult(mode="tabular", tables=tables, meta={"sheets": len(tables)})


def _ext_csv(path: Path) -> ExtractionResult:
    # Read bytes first so we can reject binary files masquerading as CSV
    # (a NUL-heavy file would otherwise be ingested as garbage text).
    raw = path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise ValueError("File contains NUL bytes; not a valid CSV/text file.")
    text = raw.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows: list[list[str]] = []
    for i, row in enumerate(reader):
        if i >= MAX_ROWS:
            log.warning("extract.csv.row_cap", path=str(path), cap=MAX_ROWS)
            break
        rows.append([c for c in row])
    return ExtractionResult(mode="tabular", tables=[(path.stem, rows)])


def _ext_text(path: Path) -> ExtractionResult:
    raw = path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise ValueError("File contains NUL bytes; not a valid text file.")
    text = raw.decode("utf-8", errors="replace")
    return ExtractionResult(mode="prose", text=_cap_text(text))


def _ext_html(path: Path) -> ExtractionResult:
    from html.parser import HTMLParser

    class _Stripper(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._buf: list[str] = []
            self._skip = 0

        def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ARG002
            if tag in {"script", "style"}:
                self._skip += 1

        def handle_endtag(self, tag: str) -> None:
            if tag in {"script", "style"} and self._skip:
                self._skip -= 1

        def handle_data(self, data: str) -> None:
            if not self._skip and data.strip():
                self._buf.append(data)

    raw = path.read_bytes()
    if b"\x00" in raw[:4096]:
        raise ValueError("File contains NUL bytes; not a valid HTML file.")
    text = raw.decode("utf-8", errors="replace")[:MAX_HTML_CHARS]
    s = _Stripper()
    s.feed(text)
    return ExtractionResult(mode="prose", text="\n".join(s._buf))


def _ext_doc(path: Path) -> ExtractionResult:
    """Legacy binary .doc is not supported.

    We accept the extension at upload (so the UI can tell the user to
    convert) but extraction fails with a clear, actionable message
    rather than a confusing KeyError deep in the pipeline.
    """
    raise ValueError(
        "Legacy .doc files are not supported. Please convert to .docx "
        "(or PDF) and re-upload."
    )


_EXTRACTORS = {
    "pdf": _ext_pdf,
    "docx": _ext_docx,
    "xlsx": _ext_xlsx,
    "csv": _ext_csv,
    "txt": _ext_text,
    "md": _ext_text,
    "html": _ext_html,
    "htm": _ext_html,
    "doc": _ext_doc,
}


def extract(path: Path) -> ExtractionResult:
    """Dispatch to the correct extractor based on file extension.

    C-3 — every result passes through :func:`sanitize_for_context`
    on the way out. The sanitization is applied to ``text`` (prose
    extractors) **and** to every cell of every table row in
    ``tables`` (tabular extractors). A cell value that contains
    the orchestrator's fence delimiters is scrubbed before it
    reaches the chunker; the LLM never sees the raw text.
    """
    ext = path.suffix.lower().lstrip(".")
    if ext not in settings.ALLOWED_UPLOAD_EXTS:
        raise ValueError(f"Unsupported file type: .{ext}")
    fn = _EXTRACTORS.get(ext)
    if fn is None:
        raise ValueError(f"No extractor registered for .{ext}")
    log.info("extract.start", path=str(path), ext=ext)
    result = fn(path)
    # C-3: sanitize the prose path. This is the common case for
    # PDF/DOCX/TXT/HTML; tabular extractors sanitize per-cell below.
    if result.text:
        result.text = sanitize_for_context(result.text)
    # C-3: sanitize every cell in the tabular path. We mutate
    # the row in place; the chunker reads from the same objects.
    if result.tables:
        sanitized_tables: list[Tuple[str, List[List[str]]]] = []
        for sheet_name, rows in result.tables:
            sanitized_rows = [
                [sanitize_for_context(str(cell)) for cell in row]
                for row in rows
            ]
            sanitized_tables.append((sheet_name, sanitized_rows))
        result.tables = sanitized_tables
    log.info(
        "extract.done",
        path=str(path),
        mode=result.mode,
        chars=len(result.text),
        tables=len(result.tables),
    )
    return result